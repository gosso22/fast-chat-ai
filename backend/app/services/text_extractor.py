"""
Text extraction service for multiple file formats.
Supports PDF, DOCX, TXT, and Markdown files with error handling and metadata preservation.
"""

import io
import logging
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Dict, Any, Optional, Union

import PyPDF2
import pdfplumber
import markdown
from docx import Document
from fastapi import UploadFile

logger = logging.getLogger(__name__)


class TextExtractionError(Exception):
    """Custom exception for text extraction errors."""
    
    def __init__(self, message: str, file_type: str, original_error: Optional[Exception] = None):
        self.message = message
        self.file_type = file_type
        self.original_error = original_error
        super().__init__(self.message)


class ExtractionResult:
    """Result of text extraction with metadata."""
    
    def __init__(
        self,
        text: str,
        metadata: Dict[str, Any],
        word_count: int,
        character_count: int,
        extraction_method: str
    ):
        self.text = text
        self.metadata = metadata
        self.word_count = word_count
        self.character_count = character_count
        self.extraction_method = extraction_method
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert result to dictionary."""
        return {
            "text": self.text,
            "metadata": self.metadata,
            "word_count": self.word_count,
            "character_count": self.character_count,
            "extraction_method": self.extraction_method
        }


class BaseTextExtractor(ABC):
    """Abstract base class for text extractors."""
    
    @abstractmethod
    def extract(self, file_content: bytes, filename: str) -> ExtractionResult:
        """Extract text from file content."""
        pass
    
    @abstractmethod
    def supports_format(self, file_extension: str) -> bool:
        """Check if extractor supports the given file format."""
        pass
    
    def _count_words_and_chars(self, text: str) -> tuple[int, int]:
        """Count words and characters in text."""
        word_count = len(text.split()) if text.strip() else 0
        char_count = len(text)
        return word_count, char_count


class PDFTextExtractor(BaseTextExtractor):
    """Text extractor for PDF files using PyPDF2 and pdfplumber as fallback."""
    
    def supports_format(self, file_extension: str) -> bool:
        """Check if this extractor supports PDF files."""
        return file_extension.lower() == '.pdf'
    
    def extract(self, file_content: bytes, filename: str) -> ExtractionResult:
        """Extract text from PDF file."""
        try:
            # First try with PyPDF2 (faster)
            result = self._extract_with_pypdf2(file_content, filename)
            if result and result.text.strip():
                return result
            
            # Fallback to pdfplumber (better for complex layouts)
            logger.info(f"PyPDF2 extraction yielded empty text for {filename}, trying pdfplumber")
            return self._extract_with_pdfplumber(file_content, filename)
            
        except Exception as e:
            raise TextExtractionError(
                f"Failed to extract text from PDF: {str(e)}",
                "pdf",
                e
            )
    
    def _extract_with_pypdf2(self, file_content: bytes, filename: str) -> ExtractionResult:
        """Extract text using PyPDF2."""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            # Check if PDF is encrypted
            if pdf_reader.is_encrypted:
                raise TextExtractionError(
                    "PDF is encrypted and cannot be processed",
                    "pdf"
                )
            
            text_parts = []
            metadata = {
                "page_count": len(pdf_reader.pages),
                "title": None,
                "author": None,
                "subject": None,
                "creator": None
            }
            
            # Extract metadata if available
            if pdf_reader.metadata:
                metadata.update({
                    "title": pdf_reader.metadata.get("/Title"),
                    "author": pdf_reader.metadata.get("/Author"),
                    "subject": pdf_reader.metadata.get("/Subject"),
                    "creator": pdf_reader.metadata.get("/Creator")
                })
            
            # Extract text from each page
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"[Page {page_num + 1}]\n{page_text}")
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                    continue
            
            full_text = "\n\n".join(text_parts)
            word_count, char_count = self._count_words_and_chars(full_text)
            
            return ExtractionResult(
                text=full_text,
                metadata=metadata,
                word_count=word_count,
                character_count=char_count,
                extraction_method="PyPDF2"
            )
            
        except Exception as e:
            if isinstance(e, TextExtractionError):
                raise
            raise TextExtractionError(
                f"PyPDF2 extraction failed: {str(e)}",
                "pdf",
                e
            )
    
    def _extract_with_pdfplumber(self, file_content: bytes, filename: str) -> ExtractionResult:
        """Extract text using pdfplumber (better for complex layouts)."""
        try:
            pdf_file = io.BytesIO(file_content)
            
            with pdfplumber.open(pdf_file) as pdf:
                text_parts = []
                metadata = {
                    "page_count": len(pdf.pages),
                    "title": pdf.metadata.get("Title") if pdf.metadata else None,
                    "author": pdf.metadata.get("Author") if pdf.metadata else None,
                    "subject": pdf.metadata.get("Subject") if pdf.metadata else None,
                    "creator": pdf.metadata.get("Creator") if pdf.metadata else None
                }
                
                # Extract text from each page
                for page_num, page in enumerate(pdf.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(f"[Page {page_num + 1}]\n{page_text}")
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                        continue
                
                full_text = "\n\n".join(text_parts)
                word_count, char_count = self._count_words_and_chars(full_text)
                
                return ExtractionResult(
                    text=full_text,
                    metadata=metadata,
                    word_count=word_count,
                    character_count=char_count,
                    extraction_method="pdfplumber"
                )
                
        except Exception as e:
            raise TextExtractionError(
                f"pdfplumber extraction failed: {str(e)}",
                "pdf",
                e
            )


class DOCXTextExtractor(BaseTextExtractor):
    """Text extractor for DOCX files using python-docx."""
    
    def supports_format(self, file_extension: str) -> bool:
        """Check if this extractor supports DOCX files."""
        return file_extension.lower() == '.docx'
    
    def extract(self, file_content: bytes, filename: str) -> ExtractionResult:
        """Extract text from DOCX file."""
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            
            text_parts = []
            metadata = {
                "paragraph_count": 0,
                "table_count": 0,
                "title": None,
                "author": None,
                "subject": None,
                "created": None,
                "modified": None
            }
            
            # Extract core properties if available
            if hasattr(doc, 'core_properties'):
                core_props = doc.core_properties
                metadata.update({
                    "title": core_props.title,
                    "author": core_props.author,
                    "subject": core_props.subject,
                    "created": str(core_props.created) if core_props.created else None,
                    "modified": str(core_props.modified) if core_props.modified else None
                })
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
                    metadata["paragraph_count"] += 1
            
            # Extract text from tables
            for table in doc.tables:
                metadata["table_count"] += 1
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(" | ".join(row_text))
                
                if table_text:
                    text_parts.append(f"[Table {metadata['table_count']}]\n" + "\n".join(table_text))
            
            full_text = "\n\n".join(text_parts)
            word_count, char_count = self._count_words_and_chars(full_text)
            
            return ExtractionResult(
                text=full_text,
                metadata=metadata,
                word_count=word_count,
                character_count=char_count,
                extraction_method="python-docx"
            )
            
        except Exception as e:
            raise TextExtractionError(
                f"Failed to extract text from DOCX: {str(e)}",
                "docx",
                e
            )


class MarkdownTextExtractor(BaseTextExtractor):
    """Text extractor for Markdown files."""
    
    def supports_format(self, file_extension: str) -> bool:
        """Check if this extractor supports Markdown files."""
        return file_extension.lower() in ['.md', '.markdown']
    
    def extract(self, file_content: bytes, filename: str) -> ExtractionResult:
        """Extract text from Markdown file."""
        try:
            # Decode bytes to string
            try:
                text = file_content.decode('utf-8')
            except UnicodeDecodeError:
                # Try other encodings
                for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                    try:
                        text = file_content.decode(encoding)
                        break
                    except UnicodeDecodeError:
                        continue
                else:
                    raise TextExtractionError(
                        "Could not decode markdown file with any supported encoding",
                        "markdown"
                    )
            
            # Parse markdown to extract metadata and structure
            md = markdown.Markdown(extensions=['meta', 'toc'])
            html_content = md.convert(text)
            
            metadata = {
                "encoding": "utf-8",
                "has_metadata": bool(getattr(md, 'Meta', {})),
                "heading_count": text.count('#'),
                "link_count": text.count('['),
                "code_block_count": text.count('```'),
                "meta": getattr(md, 'Meta', {})
            }
            
            # Keep original markdown text (preserve formatting)
            word_count, char_count = self._count_words_and_chars(text)
            
            return ExtractionResult(
                text=text,
                metadata=metadata,
                word_count=word_count,
                character_count=char_count,
                extraction_method="markdown"
            )
            
        except Exception as e:
            if isinstance(e, TextExtractionError):
                raise
            raise TextExtractionError(
                f"Failed to extract text from Markdown: {str(e)}",
                "markdown",
                e
            )


class PlainTextExtractor(BaseTextExtractor):
    """Text extractor for plain text files."""
    
    def supports_format(self, file_extension: str) -> bool:
        """Check if this extractor supports plain text files."""
        return file_extension.lower() == '.txt'
    
    def extract(self, file_content: bytes, filename: str) -> ExtractionResult:
        """Extract text from plain text file."""
        try:
            # Try to decode with UTF-8 first
            try:
                text = file_content.decode('utf-8')
                encoding = 'utf-8'
            except UnicodeDecodeError:
                # Try other common encodings
                for enc in ['latin-1', 'cp1252', 'iso-8859-1']:
                    try:
                        text = file_content.decode(enc)
                        encoding = enc
                        break
                    except UnicodeDecodeError:
                        continue
                else:
                    raise TextExtractionError(
                        "Could not decode text file with any supported encoding",
                        "txt"
                    )
            
            metadata = {
                "encoding": encoding,
                "line_count": len(text.splitlines()),
                "empty_lines": len([line for line in text.splitlines() if not line.strip()])
            }
            
            word_count, char_count = self._count_words_and_chars(text)
            
            return ExtractionResult(
                text=text,
                metadata=metadata,
                word_count=word_count,
                character_count=char_count,
                extraction_method="plain_text"
            )
            
        except Exception as e:
            if isinstance(e, TextExtractionError):
                raise
            raise TextExtractionError(
                f"Failed to extract text from plain text file: {str(e)}",
                "txt",
                e
            )


class TextExtractionService:
    """Main service for text extraction from various file formats."""
    
    def __init__(self):
        self.extractors = [
            PDFTextExtractor(),
            DOCXTextExtractor(),
            MarkdownTextExtractor(),
            PlainTextExtractor()
        ]
    
    def extract_text(self, file: Union[UploadFile, bytes], filename: str) -> ExtractionResult:
        """
        Extract text from uploaded file.
        
        Args:
            file: Either UploadFile object or bytes content
            filename: Name of the file
            
        Returns:
            ExtractionResult with text and metadata
            
        Raises:
            TextExtractionError: If extraction fails or format not supported
        """
        try:
            # Get file content as bytes
            if isinstance(file, bytes):
                file_content = file
            else:
                file.file.seek(0)
                file_content = file.file.read()
                file.file.seek(0)
            
            # Determine file extension
            file_extension = Path(filename).suffix.lower()
            
            # Find appropriate extractor
            extractor = self._get_extractor(file_extension)
            if not extractor:
                raise TextExtractionError(
                    f"No extractor available for file type: {file_extension}",
                    file_extension
                )
            
            # Extract text
            result = extractor.extract(file_content, filename)
            
            # Validate extraction result
            if not result.text or not result.text.strip():
                raise TextExtractionError(
                    f"No text content extracted from file: {filename}",
                    file_extension
                )
            
            logger.info(
                f"Successfully extracted {result.word_count} words "
                f"from {filename} using {result.extraction_method}"
            )
            
            return result
            
        except TextExtractionError:
            raise
        except Exception as e:
            raise TextExtractionError(
                f"Unexpected error during text extraction: {str(e)}",
                Path(filename).suffix.lower(),
                e
            )
    
    def _get_extractor(self, file_extension: str) -> Optional[BaseTextExtractor]:
        """Get appropriate extractor for file extension."""
        for extractor in self.extractors:
            if extractor.supports_format(file_extension):
                return extractor
        return None
    
    def get_supported_formats(self) -> list[str]:
        """Get list of supported file formats."""
        formats = []
        for extractor in self.extractors:
            if isinstance(extractor, PDFTextExtractor):
                formats.append('.pdf')
            elif isinstance(extractor, DOCXTextExtractor):
                formats.append('.docx')
            elif isinstance(extractor, MarkdownTextExtractor):
                formats.extend(['.md', '.markdown'])
            elif isinstance(extractor, PlainTextExtractor):
                formats.append('.txt')
        return formats